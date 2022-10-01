import streamlit as st

# import gradio as gr
# Def_04 Docx file to translated_Docx file
from transformers import MarianMTModel, MarianTokenizer
import nltk
from nltk.tokenize import sent_tokenize
from nltk.tokenize import LineTokenizer
nltk.download('punkt')
import math
import torch
from docx import Document
from time import sleep

import docx
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

# Def_01 applying process bar to function
import sys

def print_progress_bar(index, total, label):
    n_bar = 50  # Progress bar width
    progress = index / total
    sys.stdout.write('\r')
    sys.stdout.write(f"[{'=' * int(n_bar * progress):{n_bar}s}] {int(100 * progress)}%  {label}")
    sys.stdout.flush()
    

if torch.cuda.is_available():  
  dev = "cuda"
else:  
  dev = "cpu" 
device = torch.device(dev)
 
mname = 'Helsinki-NLP/opus-mt-en-hi'
tokenizer = MarianTokenizer.from_pretrained(mname)
model = MarianMTModel.from_pretrained(mname)
model.to(device)


def btTranslator(docxfile):
  a=getText(docxfile)
  a1=a.split('\n')
  bigtext='''  '''
  for a in a1:
    bigtext=bigtext+'\n'+a
  #files=Document()
  lt = LineTokenizer()
  batch_size = 8
  paragraphs = lt.tokenize(bigtext)   
  translated_paragraphs = []


  for index, paragraph in enumerate(paragraphs):
    # ######################################
      total=len(paragraphs)
      print_progress_bar(index, total, "Percentage Bar")
      sleep(0.5)

    # ######################################
      sentences = sent_tokenize(paragraph)
      batches = math.ceil(len(sentences) / batch_size)     
      translated = []
      for i in range(batches):
          sent_batch = sentences[i*batch_size:(i+1)*batch_size]
          model_inputs = tokenizer(sent_batch, return_tensors="pt", padding=True, truncation=True, max_length=500).to(device)
          with torch.no_grad():
              translated_batch = model.generate(**model_inputs)
          translated += translated_batch
      translated = [tokenizer.decode(t, skip_special_tokens=True) for t in translated]
      translated_paragraphs += [" ".join(translated)]
      #files.add_paragraph(translated)
  translated_text = "\n".join(translated_paragraphs)
  
  #f=files.save(f"NewFileName.docx")
  #return files.save(f"New_File.docx")
  return translated_text
st.title('Translator App')
st.markdown("Translate from Docx file")
st.sidebar.subheader("File Upload")

datas=st.sidebar.file_uploader("Original File")
# data=getText("C:\Users\Ambresh C\Desktop\Python Files\Translators\Trail Doc of 500 words.docx")

#st.sidebar.download_button(label='Download Translated File',file_name='Translated.docx', data=btTranslator(datas)) 
st.write(btTranslator(datas))
